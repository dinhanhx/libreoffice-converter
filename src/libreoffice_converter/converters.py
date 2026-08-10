import copy
import logging
import os
import re
import subprocess
import zipfile
from pathlib import Path

import anyio
from anyio.to_thread import run_sync
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

SOFFICE_TIMEOUT = int(os.getenv("SOFFICE_TIMEOUT", 300))
UNOSERVER_PORT = os.getenv("UNOSERVER_PORT", "2003")
MAX_CONCURRENT = int(os.getenv("MAX_CONCURRENT_CONVERSIONS", 2))
PREP_DOCX = os.getenv("PREP_DOCX", "false").lower() == "true"
STRIP_RSIDS = os.getenv("STRIP_RSIDS", "false").lower() == "true"
FLATTEN_MERGED_CELLS = os.getenv("FLATTEN_MERGED_CELLS", "false").lower() == "true"

_semaphore = anyio.Semaphore(MAX_CONCURRENT)

_RSID_ATTR_RE = re.compile(rb'\s*w:rsid\w*="[0-9A-Fa-f]+"')
_RSIDS_ELEMENT_RE = re.compile(rb"<w:rsids>.*?</w:rsids>", re.S)


def _strip_rsids(input_path: Path) -> None:
    tmp_path = input_path.with_suffix(input_path.suffix + ".tmp")
    with zipfile.ZipFile(input_path) as zin, zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                data = _RSID_ATTR_RE.sub(b"", data)
            elif item.filename == "word/settings.xml":
                data = _RSIDS_ELEMENT_RE.sub(b"", data)
            zout.writestr(item, data)
    tmp_path.replace(input_path)
    logger.info("prep_docx: stripped rsids from %s", input_path)


def _flatten_merged_cells(tbl) -> int:
    """Remove w:vMerge so LibreOffice doesn't have to resolve vertical cell merges.

    Some docx tables make LibreOffice's table layout hang for minutes; dropping
    vMerge (continuation cells stay as plain empty cells instead of being merged
    into the cell above) has been observed to fix this without needing to touch
    gridSpan/tcW, which this codebase's documents encode inconsistently.
    """
    changed = 0
    for tc in tbl.iter(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        vMerge = tcPr.find(qn("w:vMerge"))
        if vMerge is not None:
            tcPr.remove(vMerge)
            changed += 1
    return changed


def _strip_uniform_tc_borders(tbl) -> int:
    """Move a table's cell-level w:tcBorders up to table-level w:tblBorders.

    Deleting every cell's w:tcBorders lets LibreOffice skip per-cell border
    resolution on large tables, which is faster than iterating cells. But a
    missing w:tcBorders falls back to the table's own w:tblBorders (often
    left as "nil" by Word's exporter), not to "no override" — so cells that
    had real borders would silently go borderless. Only strip when every
    cell in the table shares the same border spec, so promoting it to
    w:tblBorders preserves the rendered result; tables with mixed per-cell
    borders (e.g. a bordered header row) are left untouched.
    """
    tcs = list(tbl.iter(qn("w:tc")))
    border_elements = []
    for tc in tcs:
        tcPr = tc.find(qn("w:tcPr"))
        border_elements.append(tcPr.find(qn("w:tcBorders")) if tcPr is not None else None)

    if not border_elements or border_elements[0] is None:
        return 0

    def same_borders(a, b) -> bool:
        if a is None or b is None:
            return False
        if [c.tag for c in a] != [c.tag for c in b]:
            return False
        return all(ca.attrib == cb.attrib for ca, cb in zip(a, b))

    shared = border_elements[0]
    if any(not same_borders(shared, b) for b in border_elements):
        return 0

    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return 0

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
    # CT_TblBorders has no tl2br/tr2bl (those are cell-only diagonal borders).
    allowed = {qn(f"w:{name}") for name in ("top", "start", "left", "bottom", "end", "right", "insideH", "insideV")}
    for child in shared:
        if child.tag in allowed:
            tblBorders.append(copy.deepcopy(child))
    tblPr.append(tblBorders)

    removed = 0
    for tc, borders in zip(tcs, border_elements):
        tc.find(qn("w:tcPr")).remove(borders)
        removed += 1
    return removed


def _prep_docx(input_path: Path) -> None:
    doc = Document(str(input_path))
    removed = 0
    for table in doc.tables:
        removed += _strip_uniform_tc_borders(table._tbl)
    doc.save(str(input_path))
    logger.info("prep_docx: removed %d tcBorders from %s", removed, input_path)

    if FLATTEN_MERGED_CELLS:
        doc = Document(str(input_path))
        flattened = 0
        for table in doc.tables:
            flattened += _flatten_merged_cells(table._tbl)
        doc.save(str(input_path))
        logger.info("prep_docx: flattened %d merged cells in %s", flattened, input_path)

    if STRIP_RSIDS:
        _strip_rsids(input_path)


async def _prep_docx_async(input_path: Path) -> None:
    if PREP_DOCX:
        logger.debug("prep_docx: starting for %s", input_path)
        await run_sync(lambda: _prep_docx(input_path))


def _run_unoconvert(input_path: Path, output_path: Path, convert_to: str) -> Path:
    logger.info("unoconvert: %s -> %s (%s)", input_path, output_path, convert_to)
    result = subprocess.run(
        [
            "unoconvert",
            "--port", UNOSERVER_PORT,
            "--convert-to", convert_to,
            str(input_path),
            str(output_path),
        ],
        capture_output=True,
        text=True,
        timeout=SOFFICE_TIMEOUT,
    )
    if result.returncode != 0:
        logger.error("unoconvert failed for %s: %s", input_path, result.stderr or result.stdout)
        raise RuntimeError(result.stderr or result.stdout)
    if not output_path.exists():
        logger.error("unoconvert did not produce expected output: %s", output_path)
        raise RuntimeError(f"unoconvert did not produce expected output: {output_path}")
    logger.info("unoconvert: produced %s", output_path)
    return output_path


async def _run_unoconvert_async(input_path: Path, output_path: Path, convert_to: str) -> Path:
    async with _semaphore:
        return await run_sync(lambda: _run_unoconvert(input_path, output_path, convert_to))


async def doc_to_docx(input_path: Path, output_dir: Path) -> Path:
    return await _run_unoconvert_async(input_path, output_dir / (input_path.stem + ".docx"), "docx")


async def docx_to_pdf(input_path: Path, output_dir: Path) -> Path:
    await _prep_docx_async(input_path)
    return await _run_unoconvert_async(input_path, output_dir / (input_path.stem + ".pdf"), "pdf")


async def docx_to_html(input_path: Path, output_dir: Path) -> Path:
    await _prep_docx_async(input_path)
    return await _run_unoconvert_async(input_path, output_dir / (input_path.stem + ".html"), "html")


async def docx_to_html_zip(input_path: Path, output_dir: Path, original_stem: str) -> Path:
    import io
    import zipfile

    await _prep_docx_async(input_path)
    html_path = await _run_unoconvert_async(input_path, output_dir / (input_path.stem + ".html"), "html")
    image_paths = list(output_dir.glob(f"{input_path.stem}_html_*.png"))

    zip_path = output_dir / f"{input_path.stem}.zip"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(html_path, arcname=f"{original_stem}.html")
        for img in image_paths:
            zf.write(img, arcname=img.name)
    zip_path.write_bytes(buf.getvalue())

    html_path.unlink(missing_ok=True)
    for img in image_paths:
        img.unlink(missing_ok=True)

    return zip_path
